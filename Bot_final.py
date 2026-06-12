import asyncio
import os
import re
import json
import time
import requests
import pandas as pd
import sys

from io import BytesIO
from datetime import datetime
from tqdm import tqdm


from playwright.async_api import async_playwright, Page, TimeoutError as PlaywrightTimeoutError

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ================= CONFIGURACIÓN =================

with open("creds.json", "r", encoding="utf-8") as archivo:
    creds = json.load(archivo)

USUARIO = creds["usuario"]
PASSWORD = creds["password"]

INPUT_FILE = "Oficios.xlsx"
OUTPUT_DIR = "Documentos"

RESULTADOS_FILE = "Resultados_acuses.xlsx"
TEMP_RESULTADOS_CSV = "temp_resultados_acuses.csv"

NUM_PESTANAS = 5
BATCH_GUARDADO = 50
MAX_INTENTOS = 3

URL_LOGIN = "https://acprod.intranet.com.mx/mbom_mx_ws/mbom_mx_web/PortalLogon"

URL_FLUJO = (
    "https://acprod.intranet.com.mx:443/boixp_mx_web/boixp_mx_web/servlet/"
    "ServletOperacionWeb?OPERACION=VGOMX012&LOCALE=es_ES&"
    "DATOS_ENTRADA.FLUJO_LANZAR=GOMXFL10090"
)


# ================= UTILIDADES =================

def dividir_lista(lista, n):
    k, m = divmod(len(lista), n)
    return (
        lista[i * k + min(i, m):(i + 1) * k + min(i + 1, m)]
        for i in range(n)
    )


async def guardar_incremental_csv(datos, archivo, lock):
    if not datos:
        return

    df = pd.DataFrame(datos, columns=["Folio", "Estado", "Ruta Docx"])

    async with lock:
        existe = os.path.exists(archivo)
        df.to_csv(
            archivo,
            mode="a",
            header=not existe,
            index=False,
            encoding="utf-8-sig"
        )


def detectar_estado_y_limpiar_subautoridad(datos: dict) -> dict:
    estados = [
        "BAJA CALIFORNIA SUR", "SAN LUIS POTOSÍ", "QUINTANA ROO",
        "BAJA CALIFORNIA", "NUEVO LEÓN", "VERACRUZ", "CHIAPAS",
        "GUANAJUATO", "MICHOACÁN", "TAMAULIPAS", "QUERÉTARO", "HIDALGO",
        "COAHUILA", "DURANGO", "CAMPECHE", "COLIMA", "JALISCO", "OAXACA",
        "PUEBLA", "SONORA", "YUCATÁN", "ZACATECAS", "TLAXCALA", "NAYARIT",
        "MÉXICO", "TABASCO", "GUERRERO", "CHIHUAHUA", "SINALOA", "MORELOS",
        "AGUASCALIENTES", "CIUDAD DE MÉXICO"
    ]

    estados_ordenados = sorted(estados, key=len, reverse=True)
    subautoridad_original = datos.get("SUBAUTORIDAD", "") or ""
    subautoridad_original = subautoridad_original.upper().strip()

    estado_encontrado = None
    subautoridad_limpia = subautoridad_original

    for estado in estados_ordenados:
        patron = r"\b" + re.escape(estado) + r"\b"

        if re.search(patron, subautoridad_original):
            estado_encontrado = estado
            subautoridad_limpia = re.sub(patron, "", subautoridad_original).strip()
            subautoridad_limpia = re.sub(r"[\W_]+$", "", subautoridad_limpia).strip()
            break

    if estado_encontrado:
        estado_formateado = estado_encontrado.title()
        estado_formateado = estado_formateado.replace(" De ", " de ").replace(" La ", " la ")
    else:
        estado_formateado = None

    subautoridad_limpia_formateada = (
        subautoridad_limpia.title()
        if subautoridad_limpia
        else subautoridad_original.title()
    )

    return {
        "ESTADO": estado_formateado,
        "SUB_AUTORIDAD_LIMPIA": subautoridad_limpia_formateada.upper()
    }


# ================= GENERACIÓN DOCUMENTO =================

def generar_documento_prorroga(datos_oficio: dict, datos_extraidos: dict):
    logo_url = "https://albertolcm.github.io/public-images/imagenes/bbva.png"

    dias_prorroga = "15 días más"
    nombre_apoderado = "Lic. _________________"
    cargo_apoderado = "apoderad_ legal"
    entidad_financiera = (
        "BBVA MEXICO, S.A., INSTITUCIÓN DE BANCA MÚLTIPLE, "
        "GRUPO FINANCIERO BBVA MEXICO"
    )

    res_autoridad = detectar_estado_y_limpiar_subautoridad(datos_extraidos)
    ciudad_estado = res_autoridad["ESTADO"] if res_autoridad["ESTADO"] else " "

    try:
        fecha_recepcion = datetime.strptime(
            datos_extraidos.get("FECHA_RECEPCION", ""),
            "%d/%m/%Y"
        )
    except Exception:
        fecha_recepcion = datetime.now()

    nombre_autoridad = f"C. {res_autoridad['SUB_AUTORIDAD_LIMPIA']}"

    meses = {
        1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
        5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
        9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
    }

    fecha_encabezado = (
        f"a {fecha_recepcion.day} de {meses[fecha_recepcion.month]} "
        f"de {fecha_recepcion.year}"
    )

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    try:
        response = requests.get(logo_url, timeout=15)

        if response.status_code == 200:
            img_data = BytesIO(response.content)
            header = section.header

            hdr_table = header.add_table(1, 1, Inches(5))
            hdr_table.autofit = False
            hdr_cell = hdr_table.cell(0, 0)

            tc = hdr_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")

            for border_name in ["top", "left", "bottom", "right"]:
                border_elm = OxmlElement(f"w:{border_name}")
                border_elm.set(qn("w:val"), "none")
                tcBorders.append(border_elm)

            tcPr.append(tcBorders)

            paragraph_header = hdr_cell.paragraphs[0]
            paragraph_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_header.add_run().add_picture(img_data, width=Inches(1.25))
            paragraph_header.paragraph_format.space_after = Pt(0)

    except Exception as e:
        doc.add_paragraph(f"<<< ERROR AL OBTENER LOGO: {e} >>>")

    datos_encabezado = [
        f"{ciudad_estado}, {fecha_encabezado}",
        f"Oficio: {datos_extraidos.get('OFICIO_AUT', '')}",
        f"Expediente: {datos_extraidos.get('EXP_AUTORIDAD', '')}",
        f"Consecutivo BBVA N°: {datos_extraidos.get('FOLIO_SUGO', '')}",
    ]

    for linea in datos_encabezado:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(linea)
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.add_run(
        f"{nombre_apoderado}, en mi carácter de {cargo_apoderado} "
        "de la Institución de Crédito denominada "
    )
    p.add_run(entidad_financiera).bold = True
    p.add_run(", ante esta autoridad respetuosamente comparezco y expongo:")

    doc.add_paragraph(
        "En atención a lo solicitado en su atento oficio de referencia, se informa que, "
        "desde el momento mismo de su petición, se procedió a iniciar las gestiones internas "
        "pertinentes para dar atención al requerimiento, sin embargo, "
        "manifestamos nuestra imposibilidad técnica y operativa para realizar "
        "lo requerido, en el plazo concedido, ya que debido a los procesos internos de "
        "disciplina, auditoria y logística, resulta insuficiente el termino para completar "
        "las gestiones, considerando que todos los informes son atendidos en la Ciudad de "
        "México"
    )

    plazo = datos_extraidos.get("PLAZO", "") or ""
    plazo_dias = plazo.split()[0] if plazo.split() else ""

    p = doc.add_paragraph()
    p.add_run(
        "Reitero la intención de mi representada para coadyuvar y cumplimentar su "
        "requerimiento, por lo que, con el único afán de dar respuesta a su petición "
        ", solicito amablemente "
    )
    p.add_run(f"{nombre_autoridad} ").bold = True
    p.add_run("se nos otorgue el termino ")
    p.add_run(f"{dias_prorroga} días hábiles").bold = True
    p.add_run(", para estar en posibilidad de atender su requerimiento correctamente.")

    p = doc.add_paragraph()
    run = p.add_run(
        "No omitimos precisar que, si las gestiones concluyen antes de lo requerido, "
        "se hará llegar de forma inmediata, lo correspondiente."
    )
    run.underline = True

    doc.add_paragraph()
    doc.add_paragraph(f"Por lo anteriormente expuesto, solicito a usted {nombre_autoridad}:")

    p = doc.add_paragraph()
    p.add_run("PRIMERO : ").bold = True
    p.add_run("Tenerme por presentado en los términos del presente escrito, en tiempo y forma.")

    p = doc.add_paragraph()
    p.add_run("SEGUNDO : ").bold = True
    p.add_run(f"Se nos otorgue el termino de {dias_prorroga} días hábiles.")

    doc.add_paragraph()

    p_protesto = doc.add_paragraph()
    p_protesto.add_run("“PROTESTO LO NECESARIO”").bold = True
    p_protesto.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_protesto.paragraph_format.space_before = Pt(12)
    p_protesto.paragraph_format.space_after = Pt(24)

    p_linea = doc.add_paragraph("__________________________")
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER

    datos_firma = [
        "BBVA MEXICO, S.A.",
        "INSTITUCIÓN DE BANCA MÚLTIPLE",
        "GRUPO FINANCIERO BBVA MEXICO",
    ]

    for line in datos_firma:
        p_data = doc.add_paragraph(line)
        p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_data.runs[0].font.size = Pt(11)
        p_data.runs[0].font.name = "Arial"
        p_data.paragraph_format.space_before = Pt(0)
        p_data.paragraph_format.space_after = Pt(0)

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    folio_doc = datos_extraidos.get("FOLIO_SUGO") or datos_oficio.get("Folio") or "SIN_FOLIO"
    nombre_archivo = f"PRORROGA {folio_doc}.docx"
    ruta_completa = os.path.join(OUTPUT_DIR, nombre_archivo)

    doc.save(ruta_completa)

    return ruta_completa


# ================= PLAYWRIGHT =================

async def obtener_estado_sesion(browser):
    context = await browser.new_context(
        ignore_https_errors=True,
        viewport={"width": 1280, "height": 800}
    )

    page = await context.new_page()

    try:
        await page.goto(URL_LOGIN, wait_until="domcontentloaded")
        await asyncio.sleep(2)
        await page.fill(".name", USUARIO)
        await page.fill(".pass", PASSWORD)
        await asyncio.sleep(2)
        async with context.expect_page() as page_info:
            await page.click(".btnEntrar")

        popup = await page_info.value
        await popup.wait_for_load_state("domcontentloaded")

        print(f"Usuario {USUARIO} autenticado")

        storage = await context.storage_state()
        await context.close()

        return storage

    except Exception as e:
        print(f"Error en login: {e}")
        await context.close()
        return None


async def obtener_valor(page: Page, label_text: str, timeout: int = 15000):
    xpath = f"//td[text()='{label_text}']/following-sibling::td[1]"

    try:
        celda = page.locator(xpath)
        await celda.wait_for(state="visible", timeout=timeout)
        texto = await celda.inner_text()
        return texto.strip()

    except PlaywrightTimeoutError:
        print(f"[AVISO] No se encontró valor para '{label_text}'")
        return None

    except Exception:
        print(f"[AVISO] No se encontró elemento para '{label_text}'")
        return None


async def cerrar_popups_extra(context, pagina_principal: Page):
    for page in context.pages:
        if page != pagina_principal:
            try:
                await page.close()
            except Exception:
                pass


async def cargar_acuse(dato: dict, page: Page, context, contador, intento: int = 1):
    folio = str(dato.get("Folio", "")).strip()

    if intento > MAX_INTENTOS:
        return [folio, "ERROR", "N/A"]

    try:
        estado = f"W{contador} | {folio}"

        await page.goto(URL_FLUJO, wait_until="domcontentloaded")
        await page.locator("#rSugo").wait_for(state="visible", timeout=15000)

        checkbox = page.locator("#rSugo")

        if not await checkbox.is_checked():
            await checkbox.click()

        await page.fill("#fSugo", folio)
        await page.click("#busqueda")

        await page.locator("#radSelec0").wait_for(state="visible", timeout=15000)

        radio = page.locator("#radSelec0")

        if not await radio.is_checked():
            await radio.click()

        await page.locator("#crono").wait_for(state="visible", timeout=15000)
        await page.click("#crono")

        await asyncio.sleep(2)

        datos_extraidos = {
            "FOLIO_SUGO": await obtener_valor(page, "FOLIO SUGO:"),
            "OFICIO_AUT": await obtener_valor(page, "OFICIO AUTORIDAD:"),
            "FECHA_RECEPCION": await obtener_valor(page, "FECHA RECEPCIÓN:"),
            "PLAZO": await obtener_valor(page, "PLAZO:"),
            "EXP_AUTORIDAD": await obtener_valor(page, "EXP.AUTORIDAD:"),
            "EXP_SUBAUT": await obtener_valor(page, "EXP.SUBAUTORIDAD:"),
            "OFICIO_SUBAUT": await obtener_valor(page, "OFICIO SUBAUT:"),
            "AUTORIDAD": await obtener_valor(page, "AUTORIDAD:"),
            "SUBAUTORIDAD": await obtener_valor(page, "SUBAUTORIDAD:"),
            "ABOGADO": await obtener_valor(page, "ABOGADO SOLICITANTE:"),
        }

        archivo_generado = generar_documento_prorroga(dato, datos_extraidos)

        return [folio, "CORRECTO", archivo_generado]

    except Exception as e:

        await cerrar_popups_extra(context, page)

        try:
            await page.reload(wait_until="domcontentloaded")
        except Exception:
            pass

        await asyncio.sleep(2)

        return await cargar_acuse(
            dato=dato,
            page=page,
            context=context,
            contador=contador,
            intento=intento + 1
        )


# ================= WORKERS PARALELOS =================

async def procesar_lote(
    id_worker,
    datos_lote,
    context,
    senal_inicio,
    pbar,
    lock_resultados
):
    buffer_resultados = []
    contador_local = 0

    if not datos_lote:
        senal_inicio.set()
        return

    page = await context.new_page()

    try:
        await page.goto(URL_FLUJO, wait_until="domcontentloaded")
        await page.locator("#busqueda").wait_for(state="visible", timeout=15000)

        senal_inicio.set()

        for idx, dato in enumerate(datos_lote, start=1):
            folio = str(dato.get("Folio", "")).strip()

            pbar.set_postfix_str(f"Worker {id_worker} | {folio}", refresh=False)

            try:
                resultado = await cargar_acuse(
                    dato=dato,
                    page=page,
                    context=context,
                    contador=f"{id_worker}-{idx}"
                )
            except Exception as e:
                resultado = [folio, "ERROR", str(e)]

            buffer_resultados.append(resultado)
            contador_local += 1

            # ÚNICO update permitido
            pbar.update(1)
            pbar.refresh()

            if contador_local % BATCH_GUARDADO == 0:
                await guardar_incremental_csv(
                    buffer_resultados,
                    TEMP_RESULTADOS_CSV,
                    lock_resultados
                )
                buffer_resultados = []

    except Exception as e:
        if not senal_inicio.is_set():
            senal_inicio.set()

    finally:
        await guardar_incremental_csv(
            buffer_resultados,
            TEMP_RESULTADOS_CSV,
            lock_resultados
        )

        try:
            await page.close()
        except Exception:
            pass

# ================= MAIN =================

async def main():
    print("\nEjecutando Bot para Prórrogas...")

    for archivo in [TEMP_RESULTADOS_CSV, RESULTADOS_FILE]:
        if os.path.exists(archivo):
            try:
                os.remove(archivo)
            except Exception:   
                pass

    try:
        datos_df = pd.read_excel(INPUT_FILE, dtype=str)
    except FileNotFoundError:
        print(f"ERROR: Archivo '{INPUT_FILE}' no encontrado.")
        return
    except Exception as e:
        print(f"ERROR leyendo Excel: {e}")
        return

    if "Folio" not in datos_df.columns:
        print("ERROR: El Excel debe tener una columna llamada 'Folio'.")
        return

    datos_df["Folio"] = datos_df["Folio"].astype(str).str.strip()
    datos_df = datos_df[datos_df["Folio"].notna()]
    datos = datos_df.to_dict(orient="records")

    print(f"Total de folios: {len(datos)}")


    time_start = time.time()

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=True,
            args=[
                "--disable-blink-features=AutomationControlled",
                "--disable-dev-shm-usage",
                "--no-sandbox"
            ],
            channel = "chrome"
        )

        storage_state = await obtener_estado_sesion(browser)

        if not storage_state:
            print("Falló el login.")
            await browser.close()
            return

        context = await browser.new_context(
            storage_state=storage_state,
            ignore_https_errors=True,
            viewport={"width": 1280, "height": 800}
        )

        lotes = list(dividir_lista(datos, NUM_PESTANAS))

        lock_resultados = asyncio.Lock()
        tasks = []

        bar_format = (
            "{l_bar}{bar}| "
            "{n_fmt}/{total_fmt} "
            "[{elapsed}<{remaining}, {rate_fmt}] "
            "{postfix}"
        )

        pbar = tqdm(
            total=len(datos),
            desc="Procesando",
            unit="folio",
            ncols=120,
            dynamic_ncols=False,
            colour="blue",
            ascii=True,
            leave=True,
            bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}, {rate_fmt}] {postfix}"
        )

        for i, lote in enumerate(lotes):
            evento_carga = asyncio.Event()

            pbar.set_description(f"Cargando Worker {i + 1}/{NUM_PESTANAS}")

            task = asyncio.create_task(
                procesar_lote(
                    id_worker=i + 1,
                    datos_lote=lote,
                    context=context,
                    senal_inicio=evento_carga,
                    pbar=pbar,
                    lock_resultados=lock_resultados
                )
            )

            tasks.append(task)
            await evento_carga.wait()

        pbar.set_description("Procesando")

        await asyncio.gather(*tasks)

        pbar.close()

        await browser.close()


    if os.path.exists(TEMP_RESULTADOS_CSV) and os.path.getsize(TEMP_RESULTADOS_CSV) > 0:
        df_resultados = pd.read_csv(
            TEMP_RESULTADOS_CSV,
            dtype=str,
            encoding="utf-8-sig"
        )

        df_resultados.to_excel(RESULTADOS_FILE, index=False)

        os.remove(TEMP_RESULTADOS_CSV)

        print(f"[OK] Resultados guardados en '{RESULTADOS_FILE}'")
        print(f"Total procesados: {len(df_resultados)}")
    else:
        print("[INFO] No se generaron resultados.")

    end_time = time.time()
    print(f"Tiempo total de ejecución: {(end_time - time_start) / 60:.2f} minutos")


if __name__ == "__main__":
    asyncio.run(main())